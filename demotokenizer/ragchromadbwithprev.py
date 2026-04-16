import tempfile
import os
import httpx
import asyncio
import time
import logging
import pytesseract
import chromadb
import duckduckgo_search
import uuid
import pandas as pd
import tkinter as tk
import re
import json

from tkinter import filedialog, messagebox, ttk
from io import BytesIO
from abc import ABC, abstractmethod
from sympy import preview
from pdf2image import convert_from_path
from langchain_community.document_loaders import PyPDFLoader
from sentence_transformers import SentenceTransformer
from PIL import Image, ImageTk
from docx import Document
from faster_whisper import WhisperModel
import sys
import shutil
import threading
from tqdm import tqdm
from dotenv import load_dotenv
from transformers import pipeline
from chromadb.config import Settings

# Load environment variables
load_dotenv()

# Set up logging to file and console
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Check dependencies at startup
def check_dependencies():
    required = ['chromadb', 'sentence_transformers', 'pytesseract', 'pdf2image', 'docx', 'faster_whisper', 'duckduckgo_search', 'httpx', 'pandas']
    for module in required:
        try:
            __import__(module)
        except ImportError:
            logger.error(f"Missing dependency: {module}")
            messagebox.showerror("Error", f"Missing dependency: {module}. Install with `pip install {module}`.")
            sys.exit(1)
    if not shutil.which("tesseract"):
        logger.error("Tesseract not found.")
        messagebox.showerror("Error", "Tesseract not found. Install Tesseract OCR.")
        sys.exit(1)
    if not shutil.which("pdftoppm"):
        logger.error("Poppler not found.")
        messagebox.showerror("Error", "Poppler not found. Install Poppler for pdf2image.")
        sys.exit(1)

check_dependencies()

# Initialize ChromaDB with configurable path
CHROMA_PATH = os.getenv("CHROMA_DB_PATH", "./chroma_db")
if not os.path.exists(CHROMA_PATH):
    os.makedirs(CHROMA_PATH)
client = chromadb.PersistentClient(path=CHROMA_PATH,settings=Settings(anonymized_telemetry=False))
collection = client.get_or_create_collection(name="document_collection", metadata={"hnsw:space": "cosine"})

# Load embedding model
embed_model = SentenceTransformer('Snowflake/arctic-embed-s')

# LaTeX cache
latex_cache = {}

# Abstract base class for document processing
class DocumentClass(ABC):
    def __init__(self, file_path):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_ext = os.path.splitext(file_path)[1].lower()

    @abstractmethod
    def extract_text(self):
        pass

class PDFDocument(DocumentClass):
    def extract_text(self):
        try:
            loader = PyPDFLoader(self.file_path)
            pages = loader.load_and_split()
            text_pages = [page.page_content for page in pages if page.page_content.strip()]
            if not text_pages:
                logger.info(f"PDF {self.file_path} has no text content; using OCR.")
                return self._extract_text_with_ocr()
            metadatas = [{"page_number": i+1} for i in range(len(text_pages))]
            logger.info(f"Extracted {len(text_pages)} pages from PDF {self.file_path}.")
            return text_pages, metadatas
        except Exception as e:
            logger.error(f"Error extracting text from PDF {self.file_path}: {e}")
            return self._extract_text_with_ocr()

    def _extract_text_with_ocr(self):
        try:
            images = convert_from_path(self.file_path)
            text_pages = [pytesseract.image_to_string(img) for img in images]
            text_pages = [page for page in text_pages if page.strip()]
            metadatas = [{"page_number": i+1} for i in range(len(text_pages))]
            logger.info(f"OCR extracted {len(text_pages)} pages from PDF {self.file_path}.")
            return text_pages, metadatas
        except Exception as e:
            logger.error(f"OCR failed for {self.file_path}: {e}")
            return [], []

class ExcelDocument(DocumentClass):
    def extract_text(self):
        try:
            xls = pd.ExcelFile(self.file_path)
            json_data = {}
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(self.file_path, sheet_name=sheet_name)
                json_data[sheet_name] = df.to_dict(orient='records')
            documents = []
            metadatas = []
            for sheet_name, rows in json_data.items():
                for row_idx, row in enumerate(rows):
                    row_text = " ".join([str(value) for value in row.values() if pd.notna(value)])
                    if row_text.strip():
                        documents.append(row_text)
                        metadatas.append({"sheet_name": sheet_name, "row_index": row_idx})
            logger.info(f"Extracted {len(documents)} text documents from Excel {self.file_path}.")
            return documents, metadatas
        except Exception as e:
            logger.error(f"Error extracting text from Excel {self.file_path}: {e}")
            return [], []

class WordDocument(DocumentClass):
    def extract_text(self):
        try:
            doc = Document(self.file_path)
            text_paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            metadatas = [{"paragraph_index": i+1} for i in range(len(text_paragraphs))]
            logger.info(f"Extracted {len(text_paragraphs)} paragraphs from DOCX {self.file_path}.")
            return text_paragraphs, metadatas
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {self.file_path}: {e}")
            return [], []

class ImageDocument(DocumentClass):
    def extract_text(self):
        try:
            img = Image.open(self.file_path)
            text = pytesseract.image_to_string(img)
            if text.strip():
                logger.info(f"OCR extracted text from image {self.file_path}.")
                return [text], [{"image_index": 1}]
            logger.info(f"No text extracted from image {self.file_path}.")
            return [], []
        except Exception as e:
            logger.error(f"Error extracting text from image {self.file_path}: {e}")
            return [], []

class AudioDocument(DocumentClass):
    def extract_text(self):
        try:
            model = WhisperModel("tiny")
            segments, _ = model.transcribe(self.file_path)
            text = " ".join([seg.text for seg in segments])
            if text.strip():
                logger.info(f"Transcribed text from audio {self.file_path}.")
                return [text], [{"audio_index": 1}]
            logger.info(f"No text transcribed from audio {self.file_path}.")
            return [], []
        except Exception as e:
            logger.error(f"Error transcribing audio {self.file_path}: {e}")
            return [], []

DOCUMENT_TYPES = {
    '.pdf': PDFDocument,
    '.xlsx': ExcelDocument,
    '.xls': ExcelDocument,
    '.docx': WordDocument,
    '.png': ImageDocument,
    '.jpg': ImageDocument,
    '.jpeg': ImageDocument,
    '.wav': AudioDocument
}

def add_documents_to_chroma(documents, collection, source_file, metadatas, batch_size=100):
    if not documents:
        return
    try:
        for i in tqdm(range(0, len(documents), batch_size), desc=f"Embedding {source_file}"):
            batch_docs = documents[i:i + batch_size]
            batch_metas = metadatas[i:i + batch_size]
            batch_ids = [f"{source_file}_{j}_{str(uuid.uuid4())[:8]}" for j in range(i, min(i + batch_size, len(documents)))]
            for metadata in batch_metas:
                metadata["source"] = source_file
                metadata["file_type"] = source_file.split('.')[-1].lower()
                metadata["chunk_index"] = metadatas.index(metadata)
                metadata["hierarchy"] = f"{source_file}/{metadata.get('sheet_name', metadata.get('page_number', metadata.get('paragraph_index', metadata.get('image_index', metadata.get('audio_index', 'unknown')))))}/{metadata.get('row_index', 'unknown')}"
            embeddings = embed_model.encode(batch_docs, show_progress_bar=False)
            collection.add(
                documents=batch_docs,
                metadatas=batch_metas,
                ids=batch_ids,
                embeddings=embeddings
            )
        logger.info(f"Added {len(documents)} documents from {source_file} to ChromaDB.")
    except Exception as e:
        logger.error(f"Error adding documents to ChromaDB: {e}")
        raise

def list_documents():
    try:
        results = collection.get(include=["documents", "metadatas"])
        documents = []
        for doc_id, doc_content, metadata in zip(results["ids"], results["documents"], results["metadatas"]):
            documents.append({
                "id": doc_id,
                "content": doc_content[:200] + "..." if len(doc_content) > 200 else doc_content,
                "source": metadata.get("source", "unknown"),
                "page_number": metadata.get("page_number", "unknown"),
                "sheet_name": metadata.get("sheet_name", "unknown"),
                "row_index": metadata.get("row_index", "unknown"),
                "paragraph_index": metadata.get("paragraph_index", "unknown"),
                "image_index": metadata.get("image_index", "unknown"),
                "audio_index": metadata.get("audio_index", "unknown"),
                "chunk_index": metadata.get("chunk_index", "unknown")
            })
        logger.info(f"Retrieved {len(documents)} documents from ChromaDB.")
        return documents
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        return []

def search_web(query, max_results=3):
    results = []
    try:
        with duckduckgo_search.DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results):
                results.append(f"{r['title']} - {r['href']} - {r['body']}")
        logger.info(f"Web search for query '{query}' returned {len(results)} results.")
        return "\n\n".join(results)
    except Exception as e:
        logger.error(f"Web search error for query '{query}': {e}")
        return ""

def smart_mcp_merge(collection, user_query, selected_doc_ids=None, use_web=True, web_first=False):
    try:
        if selected_doc_ids:
            results = collection.get(ids=selected_doc_ids, include=["documents", "metadatas"])
            top_docs = results.get("documents", [])
            metadatas = results.get("metadatas", [])
        else:
            results = collection.query(query_texts=[user_query], n_results=4)
            top_docs = results.get("documents", [[]])[0]
            metadatas = results.get("metadatas", [[]])[0]
            
        formatted_docs = []
        for doc, meta in zip(top_docs, metadatas):
            source = meta.get("source", "Unknown")
            page = meta.get("page_number", meta.get("row_index", meta.get("paragraph_index", "N/A")))
            formatted_docs.append(f"--- SOURCE: {source} (Ref: {page}) ---\n{doc}")
            
        local_context = "\n\n".join(formatted_docs)
        web_context = search_web(user_query) if use_web else ""
        full_context = f"[WEB SEARCH RESULTS]\n{web_context}\n\n[LOCAL DOCUMENT CONTEXT]\n{local_context}" if web_first else f"[LOCAL DOCUMENT CONTEXT]\n{local_context}\n\n[WEB SEARCH RESULTS]\n{web_context}"
        logger.info(f"Generated context for query '{user_query}': {full_context[:500]}...")
        return full_context, local_context
    except Exception as e:
        logger.error(f"Error in smart_mcp_merge for query '{user_query}': {e}")
        return "", ""

async def fetch_chatgpt_response(url, headers, payload, retries=3, timeout_seconds=10):
    for attempt in range(retries):
        try:
            async with httpx.AsyncClient(timeout=timeout_seconds) as client:
                async with client.stream("POST", url, headers=headers, json=payload) as response:
                    response.raise_for_status()
                    full_content = ""
                    async for line in response.aiter_lines():
                        if line.strip():  # Skip empty lines
                            try:
                                json_data = json.loads(line)
                                content = json_data.get("message", {}).get("content", "")
                                if content:
                                    full_content += content
                                if json_data.get("done", False):
                                    break
                            except json.JSONDecodeError as json_err:
                                logger.error(f"JSON parsing error for line: {line}, Error: {json_err}")
                                continue
                    logger.info(f"Raw API response: {full_content[:500]}...")  # Log first 500 chars
                    return full_content.strip()
        except httpx.HTTPStatusError as http_err:
            if 500 <= http_err.response.status_code < 600 and attempt < retries - 1:
                await asyncio.sleep(2 ** attempt)
                continue
            return f"API HTTP error: {http_err.response.status_code} - {http_err.response.text}"
        except httpx.RequestError as e:
            if attempt < retries - 1:
                await asyncio.sleep(2 ** attempt)
                continue
            logger.error(f"API request error: {str(e)}")
            return f"Error connecting to API: {str(e)}"

def validate_model_name(model_name):
    pattern = r'^[a-zA-Z0-9][a-zA-Z0-9\-_\.]*[a-zA-Z0-9]$'
    if len(model_name) > 96 or '--' in model_name or '..' in model_name:
        return False
    return bool(re.match(pattern, model_name))

def query_with_huggingface(context, query, model_name="Qwen/Qwen2-7B-Instruct", device=-1):
    if not validate_model_name(model_name.split('/')[-1]):
        return f"Error: Invalid model name '{model_name}'. Use alphanumeric chars, '-', '_', or '.', max length 96, no '--' or '..'."
    
    prompt = f"""You are a helpful assistant. Use the following context to answer the question.

Context: {context}

Question: {query}
"""
    
    try:
        generator = pipeline(
            "text-generation",
            model=model_name,
            device=device,
            framework="pt"
        )
        
        response = generator(
            prompt,
            max_length=150,
            temperature=0.7,
            num_return_sequences=1,
            do_sample=True,
            truncation=True
        )
        
        answer = response[0]["generated_text"].strip()
        if answer.startswith(prompt):
            answer = answer[len(prompt):].strip()
            
        return answer
        
    except Exception as e:
        return f"Error: {str(e)}"

def validate_file(file_path):
    max_size = 100 * 1024 * 1024  # 100MB
    if not os.path.exists(file_path):
        return False, "File does not exist."
    if os.path.getsize(file_path) > max_size:
        return False, "File size exceeds 100MB."
    return True, ""

def get_db_size(path):
    total_size = 0
    for dirpath, _, filenames in os.walk(path):
        for f in filenames:
            fp = os.path.join(dirpath, f)
            total_size += os.path.getsize(fp)
    return total_size

async def process_and_query(files, user_query, selected_doc_ids, user_url, model_name, api_key, use_web, timeout):
    if not files and not selected_doc_ids:
        return "No files uploaded or documents selected.", "", ""
    
    status_messages = []
    all_docs = []
    all_metadatas = []
    
    # Check database size
    db_size = get_db_size(CHROMA_PATH)
    max_db_size = 500 * 1024 * 1024  # 500MB
    if db_size > max_db_size:
        status_messages.append("Database size exceeds 500MB. Please reset or clear old documents.")
        return "\n".join(status_messages), "", ""
    
    for file in files:
        is_valid, msg = validate_file(file)
        if not is_valid:
            status_messages.append(f"Error with {os.path.basename(file)}: {msg}")
            continue
        
        file_ext = os.path.splitext(file)[1].lower()
        file_name = os.path.basename(file)
        if file_ext not in DOCUMENT_TYPES:
            status_messages.append(f"Unsupported file type: {file_ext} for {file_name}")
            continue
        
        doc_handler = DOCUMENT_TYPES[file_ext](file)
        docs, metadatas = doc_handler.extract_text()
        if docs:
            add_documents_to_chroma(docs, collection, file_name, metadatas)
            status_messages.append(f"Processed {file_name}: {len(docs)} chunks")
        else:
            status_messages.append(f"No text extracted from {file_name}")
        all_docs.extend(docs)
        all_metadatas.extend(metadatas)
    
    if not all_docs and not selected_doc_ids:
        return "\n".join(status_messages) or "No readable text found in uploaded files.", "", ""
    
    # Retrieve ChromaDB results and full context
    full_context, local_context = smart_mcp_merge(collection, user_query, selected_doc_ids, use_web)
    
    # Query LLM
    if ':' in user_url:
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
        payload = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": "You are a helpful assistant. Use the provided context to answer the question. If the context contains source information (e.g., 'SOURCE: file.pdf (Ref: 1)'), please cite it in your answer."},
                {"role": "user", "content": f"Context: {full_context}\n\nQuestion: {user_query}"}
            ]
        }
        answer = await fetch_chatgpt_response(user_url, headers, payload, timeout_seconds=timeout)
    else:
        answer = query_with_huggingface(full_context, user_query, model_name)
    
    return "\n".join(status_messages) or "Query executed successfully.", answer, local_context

def reset_collection():
    global collection
    start_time = time.time()
    try:
        results = collection.get(include=[])
        ids = results.get("ids", [])
        if ids:
            batch_size = 1000
            for i in range(0, len(ids), batch_size):
                batch_ids = ids[i:i + batch_size]
                collection.delete(ids=batch_ids)
                logger.info(f"Deleted batch of {len(batch_ids)} IDs")
        client.delete_collection(name="document_collection")
        collection = client.get_or_create_collection(name="document_collection", metadata={"hnsw:space": "cosine"})
        elapsed_time = time.time() - start_time
        logger.info(f"Database reset completed in {elapsed_time:.2f} seconds")
        return True, "Database reset successfully!"
    except Exception as e:
        logger.error(f"Error resetting database: {str(e)}")
        return False, f"Error resetting database: {str(e)}"

class DocumentApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("webtvmedia.net - Document Intelligence RAG")
        self.geometry("1100x850")
        self.configure(bg="#f5f7fa")
        
        self._setup_styles()
        
        self.file_paths = []
        self.recent_files = []
        self.docs = []
        self.images = []
        self.loop = asyncio.new_event_loop()
        asyncio.set_event_loop(self.loop)
        
        # Header with Logo Placeholder
        header_frame = tk.Frame(self, bg="#4a148c", height=60)
        header_frame.pack(side="top", fill="x")
        header_frame.pack_propagate(False)
        
        header_label = tk.Label(header_frame, text="WEBTVMEDIA.NET - DOCUMENT INTELLIGENCE", fg="#ffeb3b", bg="#4a148c", 
                               font=("Segoe UI", 16, "bold"), padx=20)
        header_label.pack(side="left")
        
        # Main frame with paned window
        self.paned = ttk.PanedWindow(self, orient=tk.HORIZONTAL)
        self.paned.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Left frame for controls
        self.control_frame = tk.Frame(self.paned, bg="#ffffff", relief="flat")
        self.paned.add(self.control_frame, weight=1)
        
        # Right frame for preview and logs
        self.preview_frame = tk.Frame(self.paned, bg="#ffffff", relief="flat")
        self.paned.add(self.preview_frame, weight=2)
        
        # Setup Scrollable Control Area
        self._setup_controls()
        
        # Setup Preview Area
        self._setup_preview()

        self.refresh_docs()

    def _setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam') # Use clam as a base for better customization
        
        # Colors
        WTM_PURPLE_DEEP = "#4a148c"
        WTM_PURPLE_ACTION = "#7b1fa2"
        WTM_PINK_LIGHT = "#fce4ec"
        WTM_YELLOW = "#ffeb3b"
        WTM_PINK = "#e91e63"
        BG_COLOR = "#f5f7fa"
        
        # Button Styles
        style.configure("WTM.TButton", font=("Segoe UI", 10, "bold"), padding=10, background=WTM_PURPLE_ACTION, foreground="white")
        style.map("WTM.TButton", background=[('active', "#9c27b0"), ('disabled', "#a0aec0")])
        
        style.configure("Accent.TButton", font=("Segoe UI", 10, "bold"), padding=10, background=WTM_YELLOW, foreground=WTM_PURPLE_DEEP)
        style.map("Accent.TButton", background=[('active', "#fff176")])

        # Frame Styles
        style.configure("TPanedwindow", background=BG_COLOR)
        
        # Label Styles
        style.configure("Header.TLabel", font=("Segoe UI", 11, "bold"), foreground=WTM_PURPLE_DEEP, background="white")
        style.configure("Status.TLabel", font=("Segoe UI", 9, "italic"), background="white")

    def _setup_controls(self):
        # Control frame with scrollbar
        self.canvas = tk.Canvas(self.control_frame, bg="white", highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self.control_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg="white")
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # Bind resize
        self.canvas.bind('<Configure>', lambda e: self.canvas.itemconfig(self.canvas_window, width=e.width))

        # --- Control Widgets ---
        pad_options = {'padx': 15, 'pady': 8}
        
        # Upload Section
        ttk.Label(self.scrollable_frame, text="DATA INGESTION", style="Header.TLabel").pack(**pad_options, anchor="w")
        upload_btn = ttk.Button(self.scrollable_frame, text="UPLOAD DOCUMENTS", style="WTM.TButton", command=self.upload_files)
        upload_btn.pack(fill="x", **pad_options)
        
        # Query Section
        ttk.Label(self.scrollable_frame, text="INTENT QUERY", style="Header.TLabel").pack(**pad_options, anchor="w")
        self.query_var = tk.StringVar()
        query_entry = tk.Entry(self.scrollable_frame, textvariable=self.query_var, font=("Segoe UI", 11), relief="solid", bd=1)
        query_entry.pack(fill="x", **pad_options)
        
        # Config Section
        ttk.Label(self.scrollable_frame, text="LLM CONFIGURATION", style="Header.TLabel").pack(**pad_options, anchor="w")
        
        tk.Label(self.scrollable_frame, text="API URL", bg="white", font=("Segoe UI", 9)).pack(padx=15, anchor="w")
        self.url_var = tk.StringVar(value="http://localhost:11434/api/chat")
        tk.Entry(self.scrollable_frame, textvariable=self.url_var, relief="solid", bd=1).pack(fill="x", **pad_options)
        
        tk.Label(self.scrollable_frame, text="Model Name", bg="white", font=("Segoe UI", 9)).pack(padx=15, anchor="w")
        self.model_var = tk.StringVar(value="qwen3:14b")
        tk.Entry(self.scrollable_frame, textvariable=self.model_var, relief="solid", bd=1).pack(fill="x", **pad_options)
        
        self.use_web = tk.BooleanVar()
        ttk.Checkbutton(self.scrollable_frame, text="Enable Web Search Augmentation", variable=self.use_web).pack(**pad_options, anchor="w")

        # Document Selection
        ttk.Label(self.scrollable_frame, text="KNOWLEDGE BASE", style="Header.TLabel").pack(**pad_options, anchor="w")
        self.doc_list = tk.Listbox(self.scrollable_frame, selectmode="multiple", height=8, font=("Segoe UI", 9), 
                                  relief="solid", bd=1, highlightcolor="#7b1fa2")
        self.doc_list.pack(fill="x", **pad_options)
        self.doc_list.bind('<<ListboxSelect>>', self.show_doc_preview)
        
        # Actions
        run_btn = ttk.Button(self.scrollable_frame, text="RUN ANALYSIS", style="Accent.TButton", command=self.run_query)
        run_btn.pack(fill="x", **pad_options)
        
        reset_btn = ttk.Button(self.scrollable_frame, text="RESET DATABASE", style="WTM.TButton", command=self.reset_db)
        reset_btn.pack(fill="x", **pad_options)

    def _setup_preview(self):
        # Result frame
        res_header = tk.Frame(self.preview_frame, bg="#4a148c")
        res_header.pack(fill="x")
        tk.Label(res_header, text="RETRIEVED KNOWLEDGE", fg="white", bg="#4a148c", font=("Segoe UI", 10, "bold"), pady=5).pack(side="left", padx=10)
        
        self.result_text = tk.Text(self.preview_frame, wrap="word", height=12, font=("Consolas", 10), 
                                  relief="flat", bg="#fce4ec", padx=10, pady=10)
        self.result_text.pack(fill="both", expand=True)
        
        # Answer frame
        ans_header = tk.Frame(self.preview_frame, bg="#ffeb3b")
        ans_header.pack(fill="x", pady=(10, 0))
        tk.Label(ans_header, text="AI GENERATED INSIGHTS", fg="#4a148c", bg="#ffeb3b", font=("Segoe UI", 10, "bold"), pady=5).pack(side="left", padx=10)
        
        self.answer_text = tk.Text(self.preview_frame, height=15, font=("Segoe UI", 11), 
                                  relief="flat", bg="#ffffff", padx=10, pady=10)
        self.answer_text.pack(fill="both", expand=True)
        
        self.status_var = tk.StringVar(value="Ready")
        ttk.Label(self.preview_frame, textvariable=self.status_var, style="Status.TLabel").pack(side="bottom", fill="x", padx=10)

        # Processing label
        self.processing_label = tk.Label(self.preview_frame, text="⚙ EXECUTING NEURAL PIPELINE...", fg="#e91e63", 
                                        bg="#ffffff", font=("Segoe UI", 10, "bold"))
        self.processing_label.pack(side='bottom', pady=5)
        self.processing_label.pack_forget()

    def upload_files(self):
        files = filedialog.askopenfilenames(filetypes=[("Supported files", "*.pdf *.xlsx *.xls *.docx *.png *.jpg *.jpeg *.wav")])
        self.file_paths = list(files)
        self.recent_files = [os.path.basename(f) for f in self.file_paths]
        messagebox.showinfo("Info", f"Selected {len(self.file_paths)} files.")
        self.refresh_docs()

    def clear_query(self):
        self.query_var.set("")
        self.url_var.set("http://localhost:11434/api/chat")
        self.model_var.set("qwen3:14b")
        self.key_var.set("")
        self.timeout_var.set("500")
        self.use_web.set(False)
        self.answer_text.delete(1.0, tk.END)
        self.status_var.set("")
        self.doc_list.selection_clear(0, tk.END)

    def refresh_docs(self):
        self.docs = list_documents()
        self.doc_list.delete(0, tk.END)
        display_mode = self.display_mode.get()
        for doc in self.docs:
            if display_mode == "Only Loaded Documents" and doc['source'] not in self.recent_files:
                continue
            index = doc.get('sheet_name', 'unknown') if doc.get('sheet_name', 'unknown') != 'unknown' else \
                    doc.get('page_number', 'unknown') if doc.get('page_number', 'unknown') != 'unknown' else \
                    doc.get('paragraph_index', 'unknown') if doc.get('paragraph_index', 'unknown') != 'unknown' else \
                    doc.get('image_index', 'unknown') if doc.get('image_index', 'unknown') != 'unknown' else \
                    doc.get('audio_index', 'unknown')
            display = f"{doc['id']}: {doc['source']} ({index})"
            self.doc_list.insert(tk.END, display)
        self.preview_text.delete(1.0, tk.END)

    def show_doc_preview(self, event):
        selected_indices = self.doc_list.curselection()
        self.preview_text.delete(1.0, tk.END)
        if selected_indices:
            selected_doc = self.docs[selected_indices[0]]
            self.preview_text.insert(tk.END, f"Source: {selected_doc['source']}\n\nContent:\n{selected_doc['content']}")

    def show_logs(self):
        self.log_text.delete(1.0, tk.END)
        try:
            with open('app.log', 'r') as f:
                self.log_text.insert(tk.END, f.read())
        except Exception as e:
            self.log_text.insert(tk.END, f"Error reading log file: {str(e)}")
        self.log_frame.deiconify()

    def show_result_frame(self, result):
        """
        Updates the result text area with the ChromaDB query result.
        Args:
            result: The ChromaDB query result to display.
        """
        # Enable text area for editing
        self.result_text.config(state="normal")
        self.result_text.delete(1.0, tk.END)  # Clear previous content
              
        # Insert the result into the text area
        if result:
            self.result_text.insert(tk.END, result)
        else:
            self.result_text.insert(tk.END, "No results found or an error occurred.")
        
        # Make the text area read-only again
        self.result_text.config(state="disabled")

    def run_query(self):
        query = self.query_var.get()
        url = self.url_var.get()
        model = self.model_var.get()
        key = self.key_var.get()
        try:
            timeout = float(self.timeout_var.get())
        except ValueError:
            messagebox.showwarning("Warning", "Invalid timeout value. Using default (500s).")
            timeout = 500
        use_web = self.use_web.get()
        selected_indices = self.doc_list.curselection()
        selected_doc_ids = [self.docs[i]['id'] for i in selected_indices] if selected_indices else []
        
        db_docs = list_documents()
        db_is_empty = len(db_docs) == 0
        if not query:
            messagebox.showwarning("Warning", "Please enter a query.")
            return
        if not self.file_paths and db_is_empty and not selected_doc_ids:
            messagebox.showwarning("Warning", "No documents in database or selected. Please upload at least one file or select documents.")
            return
        
        self.processing_label.pack(side='top', pady=10, fill='x')
        self.status_var.set(f"Processing {len(self.file_paths)} files...")
        query_btn = [child for child in self.scrollable_frame.children.values() if isinstance(child, tk.Button) and child['text'] == "Run Query"][0]
        query_btn.config(state='disabled')
        self.update_idletasks()
        
        def process_query_thread():
            async def process_query():
                try:
                    status, answer, local_context = await process_and_query(
                        self.file_paths, query, selected_doc_ids, url, model, key, use_web, timeout
                    )
                    self.after(0, lambda: self._update_ui(status, answer, local_context))
                except Exception as e:
                    logger.error(f"Error in process_query (query: '{query}', files: {self.file_paths}, doc_ids: {selected_doc_ids}): {e}")
                    self.after(0, lambda err=e: messagebox.showerror("Error", f"Processing failed: {str(err)}"))
                finally:
                    self.after(0, lambda: self.processing_label.pack_forget())
                    self.after(0, lambda: query_btn.config(state='normal'))
            
            self.loop.run_until_complete(process_query())
        
        threading.Thread(target=process_query_thread, daemon=True).start()

    def _update_ui(self, status, answer, local_context):
        self.status_var.set(status)
        self.show_result_frame(local_context)
        self.answer_text.delete(1.0, tk.END)
        self.images = []
        if "JSON parsing error" in answer:
            self.answer_text.insert(tk.END, "Error: The API response was invalid. Please check the API server logs or try a different query.")
        else:
            self.process_and_insert(answer)
        self.refresh_docs()
        self.file_paths = []
        self.recent_files = []

    def process_and_insert(self, text):
        pattern = r'(\\\[(.*?)\\\]|\$\$(.*?)\$\$|\\\((.*?)\\\)|\$(.*?)\$)'
        matches = list(re.finditer(pattern, text, re.DOTALL))
        pos = 0
        for match in matches:
            self.answer_text.insert(tk.END, text[pos:match.start()])
            tex = next(g for g in match.groups()[1:] if g is not None)
            tex = tex.strip()
            full_match = match.group(1)
            mode = 'display' if full_match.startswith('\\[') or full_match.startswith('$$') else 'inline'
            preview_tex = f'\\[{tex}\\]' if mode == 'display' else f'\\({tex}\\)'
            
            if preview_tex in latex_cache:
                photo = latex_cache[preview_tex]
            else:
                try:
                    buf = BytesIO()
                    preview(preview_tex, output='png', viewer='BytesIO', outputbuffer=buf, euler=False, fontsize=12)
                    buf.seek(0)
                    img = Image.open(buf)
                    max_size = (500, 200) if mode == 'display' else (200, 50)
                    img.thumbnail(max_size, Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    latex_cache[preview_tex] = photo
                except Exception as e:
                    logger.error(f"Error rendering LaTeX '{tex}': {e}")
                    self.answer_text.insert(tk.END, f"[LaTeX Error: {tex}]")
                    pos = match.end()
                    continue
            self.images.append(photo)
            if mode == 'display':
                self.answer_text.insert(tk.END, '\n')
            self.answer_text.image_create(tk.END, image=photo)
            if mode == 'display':
                self.answer_text.insert(tk.END, '\n')
            pos = match.end()
        self.answer_text.insert(tk.END, text[pos:])

    def reset_db(self):
        success, message = reset_collection()
        messagebox.showinfo("Info", message)
        self.refresh_docs()

    def destroy(self):
        if hasattr(self, 'loop') and self.loop:
            self.loop.close()
        super().destroy()

if __name__ == "__main__":
    app = DocumentApp()
    app.mainloop()